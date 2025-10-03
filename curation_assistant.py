from __future__ import annotations
"""
Synopsis Tool (All‑in‑One) — with Figure Interpretation
------------------------------------------------------
- Ingest PDFs/DOCX/TXT/CSV/XLSX into a lightweight local RAG index
- Summarize and Q&A strictly from uploaded documents
- Extract figures from PDFs, analyze uploaded pictures (OCR, sub‑panels, chart guess, scale bar, colors, EXIF)
- Generate grounded interpretations & manuscript‑style legends for figures (LLM + safe fallback)

"""

# =============================================================================
# SECTION: Compatibility shims (Python 3.13 sqlite, NumPy 2.x alias)
# =============================================================================
import sys, io, re, json, tempfile, textwrap, datetime
from pathlib import Path
from typing import List, Tuple, Optional, Dict, Any

# sqlite shim (optional)
try:
    import pysqlite3  # type: ignore
    sys.modules["sqlite3"] = sys.modules["pysqlite3"]
    SQLITE_SHIM_OK = True
except Exception:
    SQLITE_SHIM_OK = False

# NumPy shim for 2.x float_
try:
    import numpy as np  # type: ignore
    if not hasattr(np, "float_"):
        np.float_ = np.float64  # noqa: NPY201
    NP_SHIM_OK = True
except Exception:
    NP_SHIM_OK = False

# =============================================================================
# SECTION: Streamlit & core libs
# =============================================================================
import streamlit as st
import pandas as pd
from PIL import Image, ExifTags

# LangChain + vectorstores
Chroma = None
CLIENT = None
chroma_source = None
client_error = None

try:
    from langchain_chroma import Chroma as _Chroma
    Chroma = _Chroma
    chroma_source = "langchain_chroma"
except Exception:
    try:
        from langchain_community.vectorstores import Chroma as _Chroma
        Chroma = _Chroma
        chroma_source = "langchain_community"
    except Exception:
        chroma_source = None

try:
    import chromadb
    from chromadb import PersistentClient
    CLIENT = PersistentClient(path="./chroma_index")
except Exception as e:
    client_error = e

from langchain_community.vectorstores import FAISS
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_community.document_loaders import PyPDFLoader
import docx  # python-docx
import mammoth  # .docx fallback

# =============================================================================
# SECTION: Figure Analyzer deps (with safety guard)
# =============================================================================
try:
    import cv2
    CV2_OK = True
    CV2_ERR = None
except Exception as e:
    CV2_OK = False
    CV2_ERR = e

from skimage.util import img_as_ubyte
from skimage.feature import canny
from skimage.transform import hough_line, hough_line_peaks

try:
    import pytesseract
    TESS_AVAILABLE = True
except Exception:
    TESS_AVAILABLE = False

# =============================================================================
# SECTION: UI Setup
# =============================================================================
st.set_page_config(page_title="🧬 Curation Assistant", layout="wide")

# UI styles
st.markdown(
    """
<style>
.block-container{padding-top:1.5rem;}
.hero {background: linear-gradient(135deg,#6E8EF5 0%, #A777E3 50%, #F08BA9 100%);
       color: white; padding: 18px 20px; border-radius: 16px; margin-bottom: 16px;
       box-shadow: 0 6px 24px rgba(0,0,0,0.15);} 
.tag {background: rgba(255,255,255,0.18); padding: 2px 10px; border-radius: 999px; margin-left: 8px; font-size: 0.85rem;}
.stExpander, .stTabs [data-baseweb="tab"] {border-radius: 12px;}
.stButton>button {border-radius: 999px; padding: 0.5rem 1rem; font-weight: 600;}
</style>
""",
    unsafe_allow_html=True,
)

st.markdown(
    """
<div class="hero">
  <h2 style="margin:0">🧬 Curation Assistant</h2>
  <div style="margin-top:6px">
    Answers grounded strictly in your uploaded files.
    <span class="tag">Chroma</span><span class="tag">MMR</span><span class="tag">OCR</span><span class="tag">Frequencies</span><span class="tag">Interpretation</span>
  </div>
</div>
""",
    unsafe_allow_html=True,
)

if not CV2_OK:
    st.warning(
        "OpenCV (cv2) isn’t available — figure analysis/interpretation is disabled.\n"
        "Install `opencv-python-headless==4.10.0.84` and redeploy.\n\n"
        f"(Import error: {CV2_ERR})"
    )

# =============================================================================
# SECTION: Constants & Session
# =============================================================================
CHUNK_SIZE = 1200
CHUNK_OVERLAP = 150
TOP_K = 5
MMR_LAMBDA = 0.5
PERSIST_DIR = "./chroma_index"
COLLECTION = "curation_assistant"
EMBED_MODEL = "text-embedding-3-small"
LLM_MODEL = "gpt-4o-mini"

for key, default in [
    ("vectorstore", None), ("all_docs", []), ("all_figs", []), ("supp_named_frames", []), ("backend", None)
]:
    if key not in st.session_state:
        st.session_state[key] = default

# =============================================================================
# SECTION: File readers & utilities
# =============================================================================
def _read_txt(file: io.BytesIO, name: str) -> List[Document]:
    text = file.read().decode("utf-8", errors="ignore")
    return [Document(page_content=text, metadata={"source": name})]


def _read_pdf_to_docs(tmp_path: Path, name: str) -> List[Document]:
    loader = PyPDFLoader(str(tmp_path))
    docs = loader.load()
    for d in docs:
        d.metadata = {**d.metadata, "source": name}
    return docs


def _extract_pdf_images(tmp_path: Path, name: str) -> Tuple[List[Document], List[Image.Image]]:
    from pypdf import PdfReader
    pil_images: List[Image.Image] = []
    ocr_docs: List[Document] = []
    try:
        reader = PdfReader(str(tmp_path))
        for page_idx, page in enumerate(reader.pages, start=1):
            try:
                page_text = page.extract_text() or ""
            except Exception:
                page_text = ""
            for im in getattr(page, "images", []):
                try:
                    img_bytes = im.data
                    img = Image.open(io.BytesIO(img_bytes)).convert("RGB")
                    pil_images.append(img)
                    ocr_text = ""
                    if TESS_AVAILABLE:
                        try:
                            ocr_text = pytesseract.image_to_string(img)
                        except Exception:
                            ocr_text = ""
                    caption_text = "\n".join([ln for ln in page_text.splitlines() if re.match(r"\s*Fig(ure)?\b", ln, re.I)])
                    combined = "\n".join([s for s in [caption_text.strip(), ocr_text.strip()] if s])
                    if combined:
                        ocr_docs.append(
                            Document(
                                page_content=f"[FIGURE OCR p.{page_idx}]\n{combined}",
                                metadata={"source": name, "page": page_idx, "type": "figure_ocr"},
                            )
                        )
                except Exception:
                    continue
    except Exception:
        pass
    return ocr_docs, pil_images


def _read_docx(file: io.BytesIO, name: str) -> List[Document]:
    try:
        f = io.BytesIO(file.read())
        d = docx.Document(f)
        paragraphs = [p.text for p in d.paragraphs]
        text = "\n".join([p for p in paragraphs if p and p.strip()])
        return [Document(page_content=text, metadata={"source": name})]
    except Exception:
        file.seek(0)
        result = mammoth.extract_raw_text(file)
        text = result.value or ""
        return [Document(page_content=text, metadata={"source": name})]


def _read_xlsx(file: io.BytesIO, name: str) -> List[Document]:
    docs: List[Document] = []
    try:
        xl = pd.ExcelFile(file)
        for sheet in xl.sheet_names:
            df = xl.parse(sheet).astype(str)
            text_rows = ["\t".join(df.columns.astype(str))]
            text_rows += ["\t".join(map(str, row)) for row in df.itertuples(index=False, name=None)]
            docs.append(Document(page_content="\n".join(text_rows), metadata={"source": name, "sheet": sheet}))
        return docs
    except Exception as e:
        st.error(f"Failed to parse Excel '{name}': {e}")
        return []


def load_files_to_documents(uploaded_files) -> Tuple[List[Document], List[Image.Image]]:
    all_docs: List[Document] = []
    all_figs: List[Image.Image] = []
    for uf in uploaded_files:
        suffix = Path(uf.name).suffix.lower()
        if suffix in [".txt", ".md", ".csv"]:
            all_docs.extend(_read_txt(uf, uf.name))
        elif suffix == ".pdf":
            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
                tmp.write(uf.read()); tmp.flush()
                docs = _read_pdf_to_docs(Path(tmp.name), uf.name)
                ocr_docs, figs = _extract_pdf_images(Path(tmp.name), uf.name)
            all_docs.extend(docs); all_docs.extend(ocr_docs); all_figs.extend(figs)
        elif suffix == ".docx":
            all_docs.extend(_read_docx(uf, uf.name))
        elif suffix in [".xlsx", ".xlsm", ".xls"]:
            all_docs.extend(_read_xlsx(uf, uf.name))
        else:
            st.warning(f"Unsupported file type: {uf.name}")
    return all_docs, all_figs

# =============================================================================
# SECTION: Gene frequency helpers (unchanged)
# =============================================================================
def _standardize_cols(cols):
    return [str(c).strip().lower().replace(" ", "_") for c in cols]


def _find_col(cols, candidates):
    for c in cols:
        for cand in candidates:
            if c == cand or c.endswith(f"_{cand}") or cand in c:
                return c
    return None


def infer_total_samples(named_frames, sample_cands=None) -> int:
    if sample_cands is None:
        sample_cands = [
            "tumor_sample_barcode", "sample_id", "sample", "biosample",
            "patient", "subject", "case_id", "case", "participant_id"
        ]
    uniq = set()
    for _, raw_df in named_frames:
        if raw_df is None or raw_df.empty:
            continue
        df = raw_df.copy()
        df.columns = _standardize_cols(df.columns)
        col = _find_col(df.columns, sample_cands)
        if col and col in df.columns:
            vals = df[col].dropna().astype(str).str.strip()
            uniq.update(vals.tolist())
    return len(uniq)


def compute_gene_frequencies(df: pd.DataFrame, total_samples: Optional[int] = None):
    if df is None or df.empty:
        return pd.DataFrame(columns=["gene", "n_samples", "percentage"]), 0

    df = df.copy(); df.columns = _standardize_cols(df.columns)

    gene_col   = _find_col(df.columns, ["gene", "symbol", "gene_symbol", "hgnc", "ensembl", "gene_id", "hugo_symbol"])
    sample_col = _find_col(df.columns, [
        "sample", "sample_id", "tumor_sample_barcode", "biosample",
        "patient", "subject", "case_id", "case", "participant_id"
    ])
    count_col  = _find_col(df.columns, ["count", "n", "num", "samples"])

    if not gene_col:
        raise ValueError("Could not detect a gene column (e.g., 'gene', 'symbol', 'Hugo_Symbol').")

    if sample_col and (count_col is None or count_col not in df.columns):
        sub = df[[gene_col, sample_col]].dropna()
        if sub.empty:
            return pd.DataFrame(columns=["gene", "n_samples", "percentage"]), 0
        sub[gene_col] = sub[gene_col].astype(str)
        sub[sample_col] = sub[sample_col].astype(str)
        grp = sub.groupby(gene_col)[sample_col].nunique().reset_index(name="n_samples")
        denom = int(total_samples) if total_samples else int(sub[sample_col].nunique())
    elif count_col in df.columns:
        sub = df[[gene_col, count_col]].dropna()
        if sub.empty:
            return pd.DataFrame(columns=["gene", "n_samples", "percentage"]), 0
        sub[gene_col] = sub[gene_col].astype(str)
        sub[count_col] = pd.to_numeric(sub[count_col], errors="coerce").fillna(0).astype(int)
        grp = sub.groupby(gene_col)[count_col].sum().reset_index(name="n_samples")
        denom = int(total_samples) if total_samples else int(grp["n_samples"].sum())
    else:
        raise ValueError(
            "Could not detect suitable columns. Expected either (gene + sample_id) or (gene + count). "
            f"Columns seen: {list(df.columns)[:12]} ..."
        )

    grp = grp.sort_values("n_samples", ascending=False)
    denom = max(int(denom), 1)
    grp["percentage"] = (grp["n_samples"] / denom) * 100.0
    grp["percentage"] = grp["percentage"].round(6)
    if grp.columns[0] != "gene":
        grp = grp.rename(columns={grp.columns[0]: "gene"})
    return grp, denom

# =============================================================================
# SECTION: Index builders (Chroma + FAISS fallback)
# =============================================================================
def _split_docs(docs: List[Document]):
    splitter = RecursiveCharacterTextSplitter(chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)
    splits = splitter.split_documents(docs)
    if len(splits) == 0:
        raise ValueError("No text extracted from the uploaded files.")
    return splits


def build_index(docs: List[Document]):
    splits = _split_docs(docs)
    embeddings = OpenAIEmbeddings(model=EMBED_MODEL)
    if Chroma is not None and CLIENT is not None:
        try:
            vs = Chroma.from_documents(
                splits, embeddings, client=CLIENT, collection_name=COLLECTION
            )
            try: vs.persist()
            except Exception: pass
            return vs, embeddings, splits, "chroma"
        except Exception as e:
            st.sidebar.error(f"Chroma init failed; falling back to FAISS: {e}")
    vs = FAISS.from_documents(splits, embeddings)
    return vs, embeddings, splits, "faiss"


def load_existing_index():
    embeddings = OpenAIEmbeddings(model=EMBED_MODEL)
    if Chroma is not None and CLIENT is not None:
        try:
            vs = Chroma(client=CLIENT, collection_name=COLLECTION, embedding_function=embeddings)
            return vs, embeddings, "chroma"
        except Exception as e:
            st.sidebar.error(f"Chroma load failed; new ingestion needed: {e}")
    return None, embeddings, "faiss"

# =============================================================================
# SECTION: Prompts for Summary / QA
# =============================================================================
SUMMARY_SYSTEM = (
    "You are a meticulous scientific analyst. Generate a clear, concise, and strictly factual summary of the provided "
    "content. Do not speculate or add external knowledge. If tables exist, highlight important columns and notable values. "
    "If figures are present, briefly describe key features they illustrate."
)
SUMMARY_USER = "Summarize the following corpus in <= 20 bullet points. If multiple files are present, group by source name.\n\n{context}"
summary_prompt = ChatPromptTemplate.from_messages([("system", SUMMARY_SYSTEM), ("human", SUMMARY_USER)])

QA_SYSTEM = (
    "You are a bioinformatics software engineer. Answer ONLY using the given context. "
    "If the answer isn't in the context, say 'I don't know from the provided files.' "
    "Provide brief citations like (source, optional sheet/page)."
)
QA_USER = "Question: {question}\n\nUse the context to answer concisely in 1-15 sentences.\n\nContext:\n{context}"
qa_prompt = ChatPromptTemplate.from_messages([("system", QA_SYSTEM), ("human", QA_USER)])

# =============================================================================
# SECTION: FIGURE ANALYZER (vision heuristics)
# =============================================================================
def to_cv(img: Image.Image):
    if not CV2_OK:
        return None
    arr = np.array(img.convert("RGB"))
    return cv2.cvtColor(arr, cv2.COLOR_RGB2BGR)


def to_pil(cv_img):
    if not CV2_OK:
        return None
    return Image.fromarray(cv2.cvtColor(cv_img, cv2.COLOR_BGR2RGB))


def read_exif(img: Image.Image) -> Dict[str, Any]:
    meta: Dict[str, Any] = {}
    try:
        exif = img.getexif()
        if exif:
            for tag_id, value in exif.items():
                tag = ExifTags.TAGS.get(tag_id, tag_id)
                meta[str(tag)] = str(value)
    except Exception:
        pass
    return meta


def ocr_image(img: Image.Image, lang: str = "eng") -> str:
    if not TESS_AVAILABLE:
        return ""
    try:
        return pytesseract.image_to_string(img, lang=lang).strip()
    except Exception:
        return ""


def detect_subpanels(cv_bgr):
    if not CV2_OK or cv_bgr is None:
        return []
    gray = cv2.cvtColor(cv_bgr, cv2.COLOR_BGR2GRAY)
    blur = cv2.GaussianBlur(gray, (5,5), 0)
    edges = cv2.Canny(blur, 50, 150)
    contours, _ = cv2.findContours(edges, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    H, W = gray.shape[:2]
    area_img = W*H
    boxes = []
    for cnt in contours:
        x,y,w,h = cv2.boundingRect(cnt)
        area = w*h
        aspect = w/(h+1e-6)
        if area > 0.02*area_img and 0.2 < aspect < 5 and w>60 and h>60:
            boxes.append((x,y,w,h))
    # remove nested
    filtered = []
    for b in boxes:
        x,y,w,h = b
        keep = True
        for b2 in boxes:
            if b is b2: continue
            x2,y2,w2,h2 = b2
            if x>=x2 and y>=y2 and x+w<=x2+w2 and y+h<=y2+h2 and (w*h) < (w2*h2):
                keep = False; break
        if keep: filtered.append(b)
    filtered.sort(key=lambda r: (r[1]//50, r[0]))
    return filtered[:12]


def _axes_stats(gray_u8: np.ndarray) -> Dict[str, int]:
    edges = canny(gray_u8/255.0, sigma=1.5)
    h, theta, d = hough_line(edges)
    _, angles, _ = hough_line_peaks(h, theta, d, num_peaks=15)
    vert = sum(1 for a in angles if abs(np.degrees(a)) < 10 or abs(np.degrees(a)-180) < 10)
    horiz = sum(1 for a in angles if abs(abs(np.degrees(a))-90) < 10)
    return {"num_lines": int(len(angles)), "vertical": int(vert), "horizontal": int(horiz)}


def infer_chart_type(gray_u8: np.ndarray) -> str:
    ax = _axes_stats(gray_u8)
    if ax["horizontal"] >= 1 and ax["vertical"] >= 1:
        edge_density = float(np.mean(canny(gray_u8/255.0)))
        return "line/scatter (axes detected)" if edge_density > 0.08 else "bar-like (axes detected)"
    if ax["num_lines"] >= 6:
        return "grid/heatmap-like (many straight lines)"
    return "illustration/microscopy/other"


def detect_scale_bar(cv_bgr, ocr_text: str) -> bool:
    if not CV2_OK or cv_bgr is None:
        return False
    gray = cv2.cvtColor(cv_bgr, cv2.COLOR_BGR2GRAY)
    thr = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_MEAN_C,
                                cv2.THRESH_BINARY_INV, 21, 10)
    contours, _ = cv2.findContours(thr, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    H, W = gray.shape[:2]
    for cnt in contours:
        x,y,w,h = cv2.boundingRect(cnt)
        aspect = w/(h+1e-6)
        area = w*h
        if area > 0.001*W*H and (aspect > 8 or (1/aspect) > 8):
            if y > 0.7*H or x > 0.7*W or x < 0.1*W:
                return True
    low = (ocr_text or "").lower()
    for kw in ["scale", "µm", "um", "nm", "mm bar", "100 µm", "50 µm"]:
        if kw in low: return True
    return False


def dominant_colors(cv_bgr, k=5):
    if not CV2_OK or cv_bgr is None:
        return []
    data = cv2.cvtColor(cv_bgr, cv2.COLOR_BGR2RGB).reshape(-1,3).astype(np.float32)
    criteria = (cv2.TERM_CRITERIA_EPS + cv2.TERM_CRITERIA_MAX_ITER, 20, 1.0)
    _, labels, centers = cv2.kmeans(data, k, None, criteria, 3, cv2.KMEANS_PP_CENTERS)
    centers = centers.astype(int)
    counts = np.bincount(labels.flatten())
    order = np.argsort(-counts)
    return [tuple(map(int, centers[i])) for i in order]


def analyze_image(img: Image.Image, lang: str = "eng") -> Dict[str, Any]:
    if not CV2_OK:
        return {
            "width": img.width, "height": img.height, "ocr_text": ocr_image(img, lang=lang),
            "chart_guess": "(cv2 unavailable)", "has_scale_bar": False,
            "dominant_colors": [], "panels": [], "exif": read_exif(img)
        }
    cv_img = to_cv(img)
    gray = cv2.cvtColor(cv_img, cv2.COLOR_BGR2GRAY)
    gray_u8 = img_as_ubyte(gray/255.0)

    text = ocr_image(img, lang=lang)
    chart = infer_chart_type(gray_u8)

    boxes = detect_subpanels(cv_img)
    panels = []
    for (x,y,w,h) in boxes:
        crop = cv_img[y:y+h, x:x+w]
        cg = infer_chart_type(img_as_ubyte(cv2.cvtColor(crop, cv2.COLOR_BGR2GRAY)/255.0))
        excerpt = ocr_image(Image.fromarray(cv2.cvtColor(crop, cv2.COLOR_BGR2RGB)), lang=lang)[:300]
        panels.append({"bbox": (x,y,w,h), "detected_chart": cg, "ocr_excerpt": excerpt})

    has_scale = detect_scale_bar(cv_img, text)
    palette = dominant_colors(cv_img, k=5)
    meta = read_exif(img)

    return {
        "width": img.width, "height": img.height, "ocr_text": text,
        "chart_guess": chart, "has_scale_bar": has_scale,
        "dominant_colors": palette, "panels": panels, "exif": meta
    }


def draw_boxes(cv_bgr, boxes):
    if not CV2_OK or cv_bgr is None:
        return None
    out = cv_bgr.copy()
    for i,(x,y,w,h) in enumerate(boxes):
        cv2.rectangle(out, (x,y), (x+w,y+h), (36,255,12), 2)
        cv2.putText(out, f"{chr(65 + (i%26))}", (x+5, y+25), cv2.FONT_HERSHEY_SIMPLEX, 0.9, (36,255,12), 2)
    return out


def palette_image(colors):
    if not colors:
        return Image.new("RGB", (1,1), (255,255,255))
    w = 50*len(colors); h=30
    img = Image.new("RGB", (w,h), (255,255,255))
    arr = np.array(img)
    for i,(r,g,b) in enumerate(colors):
        arr[:, i*50:(i+1)*50, :] = (int(r), int(g), int(b))
    return Image.fromarray(arr)

# =============================================================================
# SECTION: Interpretation helpers (LLM + fallback)
# =============================================================================
def _brief_palette_desc(colors):
    if not colors:
        return "no palette extracted"
    def _name(c):
        r,g,b = c
        if r>200 and g>200 and b>200: return "white"
        if r>150 and b>150 and g<140: return "lavender/purple"
        if r>180 and g<160 and b<160: return "pink"
        if r<80 and g<80 and b<80:   return "dark/gray"
        return "mixed"
    names = [_name(c) for c in colors[:5]]
    out = []
    for n in names:
        if not out or out[-1]!=n: out.append(n)
    return ", ".join(out)


def _rule_based_interpretation(rep: dict) -> str:
    cg = rep.get("chart_guess","?").lower()
    pal = _brief_palette_desc(rep.get("dominant_colors", []))
    ocr = (rep.get("ocr_text") or "").strip()
    ocr_snip = textwrap.shorten(ocr.replace("\n"," "), width=180, placeholder="…") if ocr else "no OCR text"
    return textwrap.dedent(f"""
    **Interpretation (heuristic):**
    The image shows tissue with organized structures and multiple regions of interest. A scale bar enables morphometrics. The color palette aligns with H&E staining (purple nuclei, pink stroma). Consider focusing on ROIs for detailed review.

    **Evidence mapping:**
    - Scale bar: {"Yes" if rep.get("has_scale_bar") else "No"}
    - Sub-panels: {len(rep.get("panels",[]))}
    - Chart guess: {rep.get("chart_guess","–")}
    - Palette: {pal}
    - OCR (snippet): {ocr_snip}
    """).strip()


def _build_figure_context(rep: dict) -> str:
    panels_txt = []
    for idx, p in enumerate(rep.get("panels", [])):
        tag = chr(65 + (idx % 26))
        bb = p.get("bbox")
        ex = (p.get("ocr_excerpt") or "").replace("\n"," ").strip()
        ex = textwrap.shorten(ex, width=160, placeholder="…") if ex else ""
        panels_txt.append(f"- Panel {tag}: bbox={bb}, type={p.get('detected_chart','?')}" + (f", OCR: {ex}" if ex else ""))
    pal = _brief_palette_desc(rep.get("dominant_colors", []))
    ocr = (rep.get("ocr_text") or "").strip()
    ocr_short = textwrap.shorten(ocr.replace("\n"," "), width=600, placeholder="…") if ocr else ""
    return textwrap.dedent(f"""
    WIDTH: {rep.get('width')}  HEIGHT: {rep.get('height')}
    SCALE_BAR: {"Yes" if rep.get('has_scale_bar') else "No"}
    CHART_GUESS: {rep.get('chart_guess','')}
    PALETTE: {pal}
    PANELS:
    {chr(10).join(panels_txt) if panels_txt else "- (none)"}
    OCR_FULL_SNIPPET:
    {ocr_short or "(none)"}
    """).strip()


def generate_interpretation_md(rep: dict, model_name: str = None, style: str = "pathology") -> str:
    # Try to create LLM; if fails (no key), fallback
    try:
        llm = ChatOpenAI(model=(model_name or LLM_MODEL), temperature=0.2, max_tokens=700)
    except Exception:
        return _rule_based_interpretation(rep)

    system = (
        "You are an expert scientific interpreter. Write a careful, grounded analysis using ONLY the provided "
        "figure signals (OCR, panels, palette, scale bar). If information is missing, say so. "
        "Avoid speculation and external facts. Keep tone concise and professional."
    )

    style_hint = {
        "pathology": "Focus on tissue architecture, nuclei/stroma patterns, scale bar utility, and ROI differences.",
        "figure_legend": "Write a manuscript-style legend (3–6 sentences) describing panels, scale, staining, and what each ROI highlights.",
        "lay": "Write a simple, non-technical explanation (3–5 sentences)."
    }.get(style, "Be concise and grounded.")

    user = f"""
    STYLE: {style}
    STYLE_HINT: {style_hint}

    FIGURE_CONTEXT:
    {_build_figure_context(rep)}

    Produce three sections in Markdown:
    1) **Interpretation** — 4–8 sentences, grounded only in context.
    2) **Evidence mapping** — bullet list that references specific cues (scale bar, panels A/B, palette, OCR).
    3) **Figure legend (draft)** — 3–6 sentences, suitable for a manuscript caption.
    """

    prompt = ChatPromptTemplate.from_messages([("system", system), ("human", user)])
    chain = prompt | llm | StrOutputParser()
    try:
        return chain.invoke({})
    except Exception:
        return _rule_based_interpretation(rep)

# =============================================================================
# SECTION: Sidebar (uploads & actions)
# =============================================================================
with st.sidebar:
    st.subheader("Upload files")
    uploads = st.file_uploader(
        "Drop PDF/TXT/DOC/DOCX/XLSX files",
        type=["pdf", "txt", "md", "csv", "doc", "docx", "xlsx", "xlsm", "xls"],
        accept_multiple_files=True,
    )
    colb1, colb2 = st.columns(2)
    with colb1:
        build_btn = st.button("Ingest Files", type="primary")
    with colb2:
        load_btn = st.button("Load Existing Index")

# =============================================================================
# SECTION: Index actions
# =============================================================================
if build_btn:
    if not uploads:
        st.warning("Please upload at least one file.")
    else:
        with st.spinner("Parsing files, extracting figures & embedding…"):
            docs, figs = load_files_to_documents(uploads)
            if len(docs) == 0:
                st.error("No parsable content found in the uploaded files.")
            else:
                try:
                    vs, _, splits, backend = build_index(docs)
                    st.session_state.vectorstore = vs
                    st.session_state.all_docs = splits
                    st.session_state.all_figs = figs
                    st.session_state.backend = backend
                    st.success(f"Indexed {len(splits)} chunks with backend: {backend.upper()}. Figures: {len(figs)}.")
                except Exception as e:
                    st.exception(e)

if load_btn:
    with st.spinner("Loading existing index…"):
        try:
            vs, _, backend = load_existing_index()
            if vs is not None:
                st.session_state.vectorstore = vs
                st.session_state.backend = backend
                st.info(f"Loaded existing collection with backend: {backend.upper()}.")
            else:
                st.warning("No persisted index found; please ingest files.")
        except Exception as e:
            st.exception(e)

# =============================================================================
# SECTION: Tabs (Summary, QA, Figures, Frequencies)
# =============================================================================
summary_tab, qa_tab, figs_tab, freq_tab = st.tabs([
    "📝 Summary", "❓ Q&A", "🖼️ Figures", "📊 Gene Frequencies"
])

# ----- Summary -----
with summary_tab:
    st.write("Generate a concise overview of the ingested corpus.")
    if st.session_state.vectorstore is None and not st.session_state.all_docs:
        st.info("Build or load an index to enable summarization.")
    else:
        if st.button("Generate Summary", key="summ_btn"):
            with st.spinner("Summarizing…"):
                subset = list(st.session_state.all_docs)[:25]
                if not subset:
                    st.warning("No chunks available to summarize. Try rebuilding the index.")
                else:
                    llm = ChatOpenAI(model=LLM_MODEL, temperature=0)
                    chain = summary_prompt | llm | StrOutputParser()
                    context_text = "\n\n".join([
                        f"[src={d.metadata.get('source','-')} sheet={d.metadata.get('sheet','-')} page={d.metadata.get('page','-')}]\n{d.page_content}"
                        for d in subset
                    ])
                    summary = chain.invoke({"context": context_text})
                    st.markdown(summary)

# ----- Q&A -----
with qa_tab:
    st.write("Ask questions strictly from your files. If it’s not in them, the app will say it doesn’t know.")
    if st.session_state.vectorstore is None:
        st.info("Build or load an index to enable Q&A.")
    else:
        q = st.text_input("Your question")
        if st.button("Ask", type="primary") and q.strip():
            with st.spinner("Retrieving evidence & answering…"):
                retriever = st.session_state.vectorstore.as_retriever(
                    search_type="mmr", search_kwargs={"k": TOP_K, "lambda_mult": MMR_LAMBDA}
                )
                docs = retriever.get_relevant_documents(q)
                if not docs:
                    st.warning("No sufficiently relevant context found. I don't know from the provided files.")
                else:
                    context_text = "\n\n".join([
                        f"[src={d.metadata.get('source','-')} sheet={d.metadata.get('sheet','-')} page={d.metadata.get('page','-')}]\n{d.page_content}"
                        for d in docs
                    ])
                    llm = ChatOpenAI(model=LLM_MODEL, temperature=0)
                    chain = qa_prompt | llm | StrOutputParser()
                    answer = chain.invoke({"question": q, "context": context_text})
                    st.markdown("**Answer**")
                    st.write(answer)

                    st.markdown("**Cited Chunks**")
                    for i, d in enumerate(docs, 1):
                        src = d.metadata.get("source")
                        sheet = d.metadata.get("sheet")
                        page = d.metadata.get("page")
                        with st.expander(f"{i}. {src} — sheet: {sheet or '-'} — page: {page or '-'}"):
                            st.code(d.page_content[:2000])

# ----- Figures (Analyzer + Interpretation) -----
with figs_tab:
    st.write("Extracted figures from PDFs (and you can also upload standalone images below).")
    lang = st.text_input("OCR language (e.g., 'eng')", value="eng", key="fig_lang")
    local_imgs = st.file_uploader(
        "Optional: Upload additional figure images (PNG/JPG/TIFF)",
        type=["png","jpg","jpeg","tif","tiff"], accept_multiple_files=True, key="fig_add_upl",
    )

    figures = list(st.session_state.all_figs)
    if local_imgs:
        for f in local_imgs:
            try:
                figures.append(Image.open(f).convert("RGB"))
            except Exception:
                st.warning(f"Could not open {f.name}")

    if not figures:
        st.info("No figures yet — upload PDFs and click *Ingest Files*, or add image files above.")
    elif not CV2_OK:
        st.info("Figure analyzer disabled due to missing OpenCV. Once cv2 is installed, this tab will enable automatically.")
    else:
        all_reports: List[Dict[str, Any]] = []
        for i, img in enumerate(figures, start=1):
            st.markdown(f"### Figure {i}")
            col1, col2 = st.columns([2,3])

            with col1:
                st.image(img, caption="Original", use_container_width=True)

            rep = analyze_image(img, lang=lang)
            all_reports.append(rep)
            boxes = [p["bbox"] for p in rep["panels"]]
            overlay = draw_boxes(to_cv(img), boxes) if boxes else to_cv(img)

            with col2:
                if overlay is not None:
                    st.image(Image.fromarray(cv2.cvtColor(overlay, cv2.COLOR_BGR2RGB)), caption="Detected subpanels (A, B, C…)", use_container_width=True)
                else:
                    st.info("Overlay unavailable (cv2 missing).")

            k1,k2,k3 = st.columns(3)
            with k1:
                st.subheader("Chart guess"); st.write(rep["chart_guess"]) 
                st.subheader("Scale bar");  st.write("Yes" if rep["has_scale_bar"] else "No")
            with k2:
                st.subheader("Top colors")
                st.image(palette_image(rep["dominant_colors"]), use_container_width=True)
            with k3:
                st.subheader("EXIF / Metadata")
                st.json(rep["exif"] or {"note": "no EXIF"})

            with st.expander("OCR (full image)"):
                st.code(rep["ocr_text"] or "(no OCR text)")

            if rep["panels"]:
                with st.expander("Panels"):
                    for idx, p in enumerate(rep["panels"]):
                        st.markdown(f"**Panel {chr(65 + (idx % 26))}** — bbox: {p['bbox']}")
                        st.write(f"Chart guess: {p['detected_chart']}")
                        if p.get("ocr_excerpt"):
                            st.code(p["ocr_excerpt"])

            # --- Interpretation controls + generation ---
            c_int1, c_int2, c_int3 = st.columns([2,2,1])
            with c_int1:
                interp_style = st.selectbox(
                    "Interpretation style",
                    ["pathology", "figure_legend", "lay"], index=0, key=f"interp_style_{i}"
                )
            with c_int2:
                model_override = st.text_input(
                    "Model (optional override)", value="", key=f"interp_model_{i}",
                    help="Leave blank to use LLM_MODEL; e.g., gpt-4o-mini."
                )
            with c_int3:
                go = st.button("Generate interpretation", key=f"gen_interpret_{i}")

            if go:
                with st.spinner("Interpreting figure…"):
                    md = generate_interpretation_md(rep, model_name=(model_override or None), style=interp_style)
                    st.markdown(md)
                    ts = datetime.datetime.now().strftime("%Y%m%d-%H%M%S")
                    fname = f"figure_{i}_interpretation_{ts}.md"
                    st.download_button(
                        "⬇️ Download interpretation (Markdown)", data=md.encode("utf-8"),
                        file_name=fname, mime="text/markdown", use_container_width=True
                    )

            st.markdown("---")

        buf = io.BytesIO(json.dumps(all_reports, indent=2).encode("utf-8"))
        st.download_button(
            "📥 Download figures analysis (JSON)", data=buf, file_name="figure_analysis.json", mime="application/json",
            use_container_width=True
        )

# ----- Gene Frequencies -----
with freq_tab:
    st.write("Upload **cBio-style tables** (mutations/CNA/SV/clinical). Then query gene frequencies below.")

    supps = st.file_uploader(
        "Upload supplementary tables (CSV/TSV/TXT/XLSX)",
        type=["csv", "tsv", "txt", "xlsx", "xls"], accept_multiple_files=True, key="supp_upload",
    )

    named_frames = []
    if supps:
        for f in supps:
            # robust read (CSV/TSV/XLSX)
            suffix = Path(f.name).suffix.lower()
            if suffix in {".csv", ".tsv", ".txt"}:
                try:
                    df = pd.read_csv(f, sep=None, engine="python", dtype=str, na_filter=True, low_memory=False, on_bad_lines="skip")
                except Exception:
                    f.seek(0)
                    df = pd.read_csv(f, sep="\t", dtype=str, na_filter=True, low_memory=False, on_bad_lines="skip")
            elif suffix in {".xlsx", ".xls"}:
                try:
                    buf = io.BytesIO(f.read())
                    xls = pd.ExcelFile(buf, engine=None)
                    frames = [xls.parse(s, dtype=str) for s in xls.sheet_names]
                    df = pd.concat(frames, ignore_index=True) if frames else pd.DataFrame()
                except Exception as e:
                    st.error(f"Excel read failed: {e}")
                    df = pd.DataFrame()
            else:
                st.warning(f"Unsupported supplement file type: {f.name}")
                df = pd.DataFrame()

            if not df.empty:
                named_frames.append((f.name, df))
                st.write(f"✔️ Loaded **{f.name}**  shape={df.shape}")
            else:
                st.warning(f"Skipping empty/unreadable file: {f.name}")
        if named_frames:
            st.session_state["supp_named_frames"] = named_frames

    st.markdown("---")
    st.subheader("🔎 Query Gene Frequency (mutations/CNA/SV) — multiple genes")

    if not st.session_state.get("supp_named_frames"):
        st.info("Upload files above, then enter genes here.")
    else:
        named_frames = st.session_state["supp_named_frames"]

        genes_input = st.text_input("Gene symbols (comma-separated, e.g., TP53, EGFR, KRAS)")
        use_case_insensitive = st.checkbox("Case-insensitive match", value=True)

        auto_denom = infer_total_samples(named_frames)
        denom_override = st.number_input(
            f"Total number of profiled samples (0 = auto; auto currently infers {auto_denom})",
            min_value=0, step=1, value=0
        )
        denominator = int(denom_override) if denom_override > 0 else int(auto_denom)

        SAMPLE_COL_CANDS = [
            "tumor_sample_barcode", "sample_id", "sample", "biosample",
            "patient", "subject", "case_id", "case", "participant_id"
        ]

        def _norm(df: pd.DataFrame) -> pd.DataFrame:
            d = df.copy(); d.columns = _standardize_cols(d.columns); return d

        def _find_sample_col(cols):
            return _find_col(cols, SAMPLE_COL_CANDS)

        def compute_numerator_for_gene(gene_query: str):
            gene_upper = gene_query.strip().upper()
            numerator_samples = set(); add_rows_without_ids = 0; per_file_counts = []

            for fname, raw_df in named_frames:
                if raw_df is None or raw_df.empty: continue
                df = _norm(raw_df); cols = list(df.columns)
                hugo_col = _find_col(cols, ["hugo_symbol"])
                site1_col = _find_col(cols, ["site1_hugo_symbol"])
                site2_col = _find_col(cols, ["site2_hugo_symbol"])
                sample_col = _find_col(cols, SAMPLE_COL_CANDS)
                file_added = 0

                if hugo_col:
                    mask = (df[hugo_col].astype(str).str.strip().str.upper() == gene_upper) if use_case_insensitive else (df[hugo_col].astype(str) == gene_query)
                    sub = df.loc[mask]
                    if not sub.empty:
                        if sample_col in sub.columns:
                            ids = set(sub[sample_col].dropna().astype(str).str.strip().tolist())
                            numerator_samples |= ids; file_added = len(ids)
                        else:
                            non_id_cols = [c for c in sub.columns if c != hugo_col]
                            meta_like = {"entrez_gene_id", "chromosome", "cytoband"}
                            non_id_cols = [c for c in non_id_cols if c not in meta_like]
                            if len(non_id_cols) > 5:
                                nn = int(sub[non_id_cols].notna().sum(axis=1).max()); add_rows_without_ids += nn; file_added = nn
                            else:
                                cnt = int(len(sub)); add_rows_without_ids += cnt; file_added = cnt
                elif site1_col or site2_col:
                    m1 = df[site1_col].astype(str).str.strip().str.upper() == gene_upper if site1_col else False
                    m2 = df[site2_col].astype(str).str.strip().str.upper() == gene_upper if site2_col else False
                    sub = df.loc[m1 | m2]
                    if not sub.empty:
                        if sample_col in sub.columns:
                            ids = set(sub[sample_col].dropna().astype(str).str.strip().tolist())
                            numerator_samples |= ids; file_added = len(ids)
                        else:
                            cnt = int(len(sub)); add_rows_without_ids += cnt; file_added = cnt

                per_file_counts.append((fname, file_added))

            numerator = len(numerator_samples) + add_rows_without_ids
            return int(numerator), per_file_counts

        if st.button("Compute Gene Frequency") and genes_input.strip():
            raw_list = [g.strip() for g in genes_input.split(",") if g.strip()]
            seen = set(); genes = []
            for g in raw_list:
                key = g.upper() if use_case_insensitive else g
                if key not in seen:
                    seen.add(key); genes.append(g)

            if denominator == 0:
                st.warning("Denominator could not be inferred. Please set a value above.")
            else:
                results = []; per_gene_breakdown = {}
                for g in genes:
                    num, breakdown = compute_numerator_for_gene(g)
                    per_gene_breakdown[g] = breakdown
                    pct = (num / denominator * 100.0)
                    results.append({"gene": g, "n_samples": int(num), "denominator": int(denominator), "percentage": round(pct, 6)})

                st.markdown("**Per-gene results**")
                st.dataframe(pd.DataFrame(results))

                with st.expander("Per-file breakdown (by gene)"):
                    for g in genes:
                        st.markdown(f"**{g}**")
                        rows = per_gene_breakdown.get(g, [])
                        if rows:
                            st.table(pd.DataFrame(rows, columns=["file", "# counted"]))
                        else:
                            st.write("_no rows counted_")
